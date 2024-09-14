import requests
import os
import pandas as pd
import docx
from typing import Final, List, Dict, Optional
from telegram import Update , InlineKeyboardMarkup, InlineKeyboardButton
from telegram.ext import Application, CommandHandler, MessageHandler, filters, ContextTypes, ConversationHandler, CallbackQueryHandler
from docx.shared import Pt
from dotenv import load_dotenv
import tempfile
from io import StringIO
from telegram import InputFile

load_dotenv()
bot_token = os.environ.get('BOT_TOKEN')
google_books_api_key = os.environ.get('GOOGLE_BOOKS_API_KEY')



GENRE, TITLE, BOOK_NAME = range(3)
readinglist = []
doc = docx.Document()

async def start_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text("Hey! Welcome to the Books Bot!")

async def books_command(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    await update.message.reply_text("Enter the genre of the book:")
    return GENRE

async def genre_handler(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    genre = update.message.text.strip()
    if not genre:
        await update.message.reply_text("Please enter a valid genre.")
        return GENRE

    try:
        books = fetch_books_by_genre(genre)
        if not books:
            await update.message.reply_text('No books found for this genre.')
            return ConversationHandler.END

        df = pd.DataFrame(books)
        csv_data = df.to_csv(index=False)
        file = StringIO(csv_data)
        file.seek(0)  
        await update.message.reply_document(document=InputFile(file, filename='books.csv'))
    except Exception as e:
        await update.message.reply_text(f'An error occurred: {e}')
    
    return ConversationHandler.END

async def preview_command(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    await update.message.reply_text("Enter the book title for which you need a preview link:")
    return TITLE

async def preview_handler(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    title = update.message.text.strip()
    if not title:
        await update.message.reply_text("Please provide a valid title.")
        return TITLE
    
    preview_link = await get_book_preview_link(title)
    if not preview_link:
        await update.message.reply_text("No preview available for this book.")
    else:
        await update.message.reply_text(f"Preview link for '{title}': {preview_link}")

    return ConversationHandler.END

async def help_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text("Here are the available commands:\n\n"
    "/start - Welcome the user.\n"
    "/books - Prompts user to type in the genre of the book.\n"
    "/preview - Get a preview link for a specific book.\n"
    "/list - Enter a book name and then use the /reading_list command to manage your list.\n"
    "/reading_list - Displays buttons to add, delete, or view the reading list.\n"
    "/help - Display this help message." 
    )

async def list_command(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    await update.message.reply_text("Enter the name of the book you want to add to the reading list:")
    return BOOK_NAME

async def handle_book_name(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:

    book_name = update.message.text.strip()

    if not book_name:

        await update.message.reply_text("Please enter a valid book name.")

        return BOOK_NAME

    

    readinglist.append({'title': book_name, 'link': ''})  

    await update.message.reply_text(f'You entered "{book_name}".\n'

                                    'To manage your reading list, use the /reading_list command.')

    await update_document()  
    return ConversationHandler.END      
    
    readinglist.append({'title': book_name, 'link': ''})  
    await update.message.reply_text(f'You entered "{book_name}".\n'
                                    'To manage your reading list, use the /reading_list command.')
    await update_document()  
    return ConversationHandler.END

async def reading_list_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    keyboard = [
        [InlineKeyboardButton("Add a book", callback_data='add_book'),
         InlineKeyboardButton("Delete a book", callback_data='delete_book')],
        [InlineKeyboardButton("View Reading List", callback_data='view_list')]
    ]

    reply_markup = InlineKeyboardMarkup(keyboard)
    await update.message.reply_text('Reading List Menu', reply_markup=reply_markup)

async def button_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    data = query.data

    if data == 'add_book':
        await query.message.reply_text('Enter the book title to add:')
        context.user_data['action'] = 'add'
    elif data == 'delete_book':
        await query.message.reply_text('Enter the book title to delete:')
        context.user_data['action'] = 'delete'
    elif data == 'view_list':
        await view_list(query, context)

async def add_book_handler(update: Update, context: ContextTypes.DEFAULT_TYPE):
    action = context.user_data.get('action')
    if action == 'add':
        book_title = update.message.text.strip()
        if book_title:
            readinglist.append({'title': book_title, 'link': ''})  
            await update.message.reply_text(f'Added "{book_title}" to the reading list!')
            
        else:
            await update.message.reply_text("Please enter a valid book title.")

async def delete_book_handler(update: Update, context: ContextTypes.DEFAULT_TYPE):
    book_title = update.message.text.strip()
    for book in readinglist:
        if book['title'] == book_title:
            readinglist.remove(book)
            await update.message.reply_text(f'Removed "{book_title}" from the reading list!')
            return
    await update.message.reply_text(f'"{book_title}" not found in the reading list!')

async def view_list(query: Update.callback_query, context: ContextTypes.DEFAULT_TYPE):
    text = "Reading List:\n"

    for book in readinglist:

        text += f"- {book['title']}\n"

    await query.message.reply_text(text)
    await update_document()
    doc_path = '/tmp/reading_list.docx'
    try:
        doc.save(doc_path)
        with open(doc_path, 'rb') as file:
            await query.message.reply_document(file)
    except Exception as e:
        await query.message.reply_text(f"Error generating document: {e}")
    finally:
        if os.path.exists(doc_path):
            os.remove(doc_path)  
async def update_document():
    doc.add_heading('Reading List', level=1)
    
    if not readinglist:
        doc.add_paragraph("The reading list is currently empty.")
    else:
        for book in readinglist:
            paragraph = doc.add_paragraph()
            title_run = paragraph.add_run(book['title'])
            title_run.bold = True
            title_run.font.size = Pt(14)
            
            if book['link']:
                link_run = paragraph.add_run(f' - Preview Link: {book["link"]}')
                link_run.font.size = Pt(12)

async def get_book_preview_link(title: str) -> Optional[str]:
    url = f'https://www.googleapis.com/books/v1/volumes?q=intitle:{title}&key={GOOGLE_BOOKS_API_KEY}&maxResults=1'
    try:
        response = requests.get(url)
        response.raise_for_status()
        data = response.json()
    except requests.RequestException as e:
        print(f"Error: {e}")
        return None

    if 'items' in data and len(data['items']) > 0:
        preview_link = data['items'][0]['volumeInfo'].get('previewLink')
        return preview_link
    return None

async def google_book(genre: str) -> List[Dict[str, str]]:
    url = f'https://www.googleapis.com/books/v1/volumes?q=subject:{genre}&key={GOOGLE_BOOKS_API_KEY}&maxResults=10'
    try:
        response = requests.get(url)
        response.raise_for_status()
        data = response.json()
    except requests.RequestException as e:
        print(f"Error: {e}")
        return []

    books = parse_google_books_response(data)
    return books

def fetch_books_by_genre(genre: str):
    url = f'https://www.googleapis.com/books/v1/volumes?q=subject:{genre}&maxResults=40'
    response = requests.get(url)
    response.raise_for_status()
    data = response.json()
    books = []
    for item in data.get('items', []):
        volume_info = item['volumeInfo']
        title = volume_info.get('title', 'Unknown Title')
        authors = ', '.join(volume_info.get('authors', ['Unknown Author']))
        books.append({'Title': title, 'Authors': authors})
    return books

def main() -> None:
    application = Application.builder().token().build()

    application.add_handler(CommandHandler("start", start_command))
    application.add_handler(CommandHandler("books", books_command))
    application.add_handler(CommandHandler("preview", preview_command))
    application.add_handler(CommandHandler("reading_list", reading_list_command))
    application.add_handler(CommandHandler("help", help_command))
    
    application.add_handler(CallbackQueryHandler(button_callback))
    
    conv_handler = ConversationHandler(
        entry_points=[CommandHandler("books", books_command), CommandHandler("list", list_command)],
        states={
            GENRE: [MessageHandler(filters.TEXT & ~filters.COMMAND, genre_handler)],
            TITLE: [MessageHandler(filters.TEXT & ~filters.COMMAND, preview_handler)],
            BOOK_NAME: [MessageHandler(filters.TEXT & ~filters.COMMAND, handle_book_name)]
        },
        fallbacks=[CommandHandler("cancel", start_command)]
    )
    application.add_handler(conv_handler)

    application.run_polling()

if __name__ == '__main__':
    main()